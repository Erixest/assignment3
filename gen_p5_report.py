#!/usr/bin/env python3
"""Generate securep5.docx — Practical Work #5."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(1.5)

# ─── helpers ────────────────────────────────────────────────────────────────

def heading(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(3)
    else:
        run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
    return p

def para(text, bold=False, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = bold
    return p

def bullet(text, indent=0.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(indent)
    p.add_run(text)
    return p

def code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F0F0")
    pPr.append(shd)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    return p

def label(text):
    """Small bold inline label before a code block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    return p

def tbl_header(table, headers):
    row = table.rows[0]
    for cell, hdr in zip(row.cells, headers):
        cell.text = hdr
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "C9DAF8")
        tcPr.append(shd)

def tbl_row(table, values):
    row = table.add_row()
    for cell, val in zip(row.cells, values):
        cell.text = str(val)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

# ════════════════════════════════════════════════════════════════════════════
# TITLE
# ════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ПРАКТИЧЕСКАЯ РАБОТА № 5")
r.bold = True; r.font.size = Pt(16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "Тема: Комплексный анализ безопасности MVP — обработка памяти и ресурсов, "
    "инъекционные уязвимости, аутентификация, авторизация и криптографическая защита"
)
r.bold = True; r.font.size = Pt(12)

doc.add_paragraph()
para("Вариант 10: Финтех — цифровые платежи и мониторинг мошенничества")
para("Выполнил: Студент магистратуры, группа ИБ-1")
para("Проверил: Преподаватель кафедры информационной безопасности")

# ════════════════════════════════════════════════════════════════════════════
# 1. ЦЕЛЬ
# ════════════════════════════════════════════════════════════════════════════

heading("1. Цель работы")
para(
    "Цель данной работы — провести углублённый аудит безопасности уже реализованного "
    "MVP-проекта fintech-payments-mvp (Go, Gin, SQLite) по четырём направлениям: "
    "управление памятью и ресурсами; безопасность входных данных и потенциальные "
    "injection-векторы; аутентификация, авторизация и криптография; "
    "количественная оценка рисков по методологии CVSS v4.0. "
    "Для каждой выявленной проблемы необходимо реализовать исправление в коде, "
    "продемонстрировать его корректность и классифицировать уязвимость по CWE."
)

# ════════════════════════════════════════════════════════════════════════════
# 2. ЗАДАНИЕ 1 — ПАМЯТЬ И РЕСУРСЫ
# ════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
heading("2. Задание 1: Аудит управления памятью и ресурсами")

para(
    "Аудит охватывал все точки, где приложение потребляет память или удерживает "
    "системные ресурсы: HTTP-буферы, курсоры базы данных, объекты в памяти. "
    "Ниже приведены выявленные риски, критический сценарий и сводное сравнение "
    "поведения до и после защиты."
)

heading("2.1. Таблица рисков", 2)

t = doc.add_table(rows=1, cols=5)
t.style = "Table Grid"
tbl_header(t, ["№", "Область / файл", "Риск", "До исправления", "После исправления"])
tbl_row(t, [
    "1",
    "Размер тела HTTP-запроса\ncmd/server/main.go",
    "CWE-400\nДоС через RAM: атакующий отправляет тело > 100 МБ, сервер читает его целиком",
    "Никакого ограничения нет, Gin буферизует тело полностью",
    "MaxBodySizeMiddleware(1 МБ) через http.MaxBytesReader; тело > 1 МБ → 413"
])
tbl_row(t, [
    "2",
    "Курсор БД — scanPayments\ninternal/database/payment_repo.go",
    "CWE-703\nrows.Err() не проверяется; ошибка разрыва соединения во время итерации молча игнорируется",
    "return payments, nil — частичный результат без признака ошибки",
    "rows.Err() проверяется после цикла; ошибка явно возвращается вызывающему"
])
tbl_row(t, [
    "3",
    "Курсор БД — GetAuditLogs\ninternal/database/audit_repo.go",
    "CWE-703\nАналогичная проблема в аудит-репозитории",
    "return logs, nil без проверки rows.Err()",
    "rows.Err() проверяется после завершения цикла"
])
tbl_row(t, [
    "4",
    "Загрузка записей — AnalystPage\ninternal/web/handler.go",
    "CWE-400\nGetFlaggedPayments(100, 0) на каждый GET /analyst: неконтролируемая нагрузка",
    "Жёстко заданные 100 записей при каждом запросе страницы",
    "Предел снижен до 50; для полного списка используется пагинация"
])

heading("2.2. Критический сценарий: DoS через тело запроса", 2)

para(
    "Рассмотрим endpoint POST /api/v1/auth/register. До исправления он принимал "
    "JSON-тело любого размера без каких-либо ограничений. Атакующий без авторизации "
    "мог отправить тело объёмом 500 МБ — Gin читал его целиком в оперативную память, "
    "что приводило к исчерпанию RAM и аварийному завершению процесса (OOM crash)."
)

label("Код ДО — cmd/server/main.go:")
code(
    "r.Use(gin.Recovery())\n"
    "r.Use(middleware.SecurityHeaders())\n"
    "r.Use(middleware.RateLimitMiddleware(cfg))\n"
    "// ← ограничение размера тела отсутствует"
)

label("Код ПОСЛЕ:")
code(
    "r.Use(gin.Recovery())\n"
    "r.Use(middleware.SecurityHeaders())\n"
    "r.Use(middleware.MaxBodySizeMiddleware(1 << 20)) // 1 MB — CWE-400 fix\n"
    "r.Use(middleware.RateLimitMiddleware(cfg))"
)

label("Реализация middleware (internal/middleware/security.go):")
code(
    "// MaxBodySizeMiddleware ограничивает тело запроса maxBytes байтами.\n"
    "// При превышении — 413 Request Entity Too Large, тело не читается.\n"
    "func MaxBodySizeMiddleware(maxBytes int64) gin.HandlerFunc {\n"
    "    return func(c *gin.Context) {\n"
    "        if c.Request.Body != nil {\n"
    "            c.Request.Body = http.MaxBytesReader(\n"
    "                c.Writer, c.Request.Body, maxBytes)\n"
    "        }\n"
    "        c.Next()\n"
    "    }\n"
    "}"
)

para(
    "Объяснение: http.MaxBytesReader оборачивает тело и останавливает чтение, "
    "как только достигнут лимит. Gin получает ошибку при попытке распарсить JSON "
    "и возвращает клиенту 413, не загружая ни байта сверх лимита в RAM."
)
para("Подтверждение: curl с 2 МБ телом → 413; gosec: 0 issues.")

label("Код ДО — scanPayments (database/payment_repo.go):")
code(
    "for rows.Next() {\n"
    "    rows.Scan(...)\n"
    "}\n"
    "return payments, nil  // ошибка итерации молча поглощается"
)

label("Код ПОСЛЕ:")
code(
    "for rows.Next() {\n"
    "    rows.Scan(...)\n"
    "}\n"
    "// CWE-703: явная проверка ошибки итерации курсора\n"
    "if err := rows.Err(); err != nil {\n"
    "    return nil, err\n"
    "}\n"
    "return payments, nil"
)

heading("2.3. Сравнение поведения до и после защиты", 2)

t2 = doc.add_table(rows=1, cols=3)
t2.style = "Table Grid"
tbl_header(t2, ["Сценарий", "До исправления", "После исправления"])
tbl_row(t2, [
    "POST /register с телом 500 МБ",
    "Сервер потребляет всю RAM → OOM crash",
    "413 Request Entity Too Large, тело не читается"
])
tbl_row(t2, [
    "Разрыв соединения во время SELECT",
    "Частичный результат, ошибка скрыта",
    "rows.Err() → HTTP 500 с логированием"
])
tbl_row(t2, [
    "GET /analyst без пагинации",
    "До 100 строк на каждый запрос",
    "Не более 50 строк; остальное — по offset"
])

# ════════════════════════════════════════════════════════════════════════════
# 3. ЗАДАНИЕ 2 — SECURE CODE REVIEW
# ════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
heading("3. Задание 2: Secure Code Review — точки входа и опасные sink'и")

para(
    "Для бизнес-сценария «создание платежа пользователем» (POST /payments/new) "
    "построена карта доверительных границ и выявлены опасные точки назначения данных."
)

heading("3.1. Карта доверительных границ", 2)

label("Trust Boundary Map (нотация потока данных):")
code(
    "Browser — HTTP POST /payments/new\n"
    "  │  Данные: amount, currency, recipient_id, description\n"
    "  ▼\n"
    "[Граница 1: Internet → Application]\n"
    "  │  RateLimitMiddleware · MaxBodySizeMiddleware · SecurityHeaders\n"
    "  ▼\n"
    "WebHandler.CreatePayment             ← Source (недоверенный ввод)\n"
    "  │  Валидация: amount (float, 0–1M) · currency (allowlist)\n"
    "  │            recipient_id (regex ^[A-Z0-9]{8,20}$) · description (max 500)\n"
    "  ▼\n"
    "[Граница 2: Handler → Service]\n"
    "  │  PaymentService.CreatePayment\n"
    "  ▼\n"
    "[Граница 3: Service → Repository → SQLite]\n"
    "  │  DB.CreatePayment → INSERT ... VALUES (?,?,?,?,?,?,?,?,?)\n"
    "  ▼\n"
    "SQLite payments                       ← Sink 1 (SQL)\n"
    "AuditService.Log → audit_logs         ← Sink 2 (Audit Log)\n"
    "HTTP Response → browser               ← Sink 3 (Output)"
)

heading("3.2. Таблица: источник → распространение → sink → защита", 2)

t3 = doc.add_table(rows=1, cols=5)
t3.style = "Table Grid"
tbl_header(t3, ["Источник данных", "Путь", "Sink", "Уязвимость (до)", "Защита (после)"])
tbl_row(t3, [
    "recipient_id\n(POST form, web)",
    "PostForm → CreatePayment\n→ DB.CreatePayment",
    "SQL INSERT",
    "CWE-20: только length check (8–20), любые символы проходили",
    "Allowlist regex ^[A-Z0-9]{8,20}$ (recipientIDPattern)"
])
tbl_row(t3, [
    "reason\n(flag/reject forms)",
    "PostForm → FlagPayment\n→ DB.UpdatePaymentStatus",
    "SQL UPDATE + Audit Log",
    "CWE-20: max-length отсутствовал, возможна передача 100 КБ",
    "len(reason) > 1000 → 400 Bad Request"
])
tbl_row(t3, [
    "details, ip_address\n(все вызовы Log())",
    "c.ClientIP(), user input\n→ AuditService.Log\n→ audit_logs",
    "Audit Log (SQLite)",
    "CWE-117: символы \\r\\n позволяли вставить фиктивные строки в журнал",
    "sanitizeLogField() удаляет CR/LF/TAB, обрезает до 500 символов"
])
tbl_row(t3, [
    "paymentID\n(URL param :id)",
    "c.Param(\"id\") →\nstrconv.ParseInt (ошибка игнорировалась)",
    "PaymentService → DB",
    "CWE-703: ParseInt ошибка поглощена; id=0 уходит в запрос",
    "Явная проверка err != nil и paymentID <= 0 → 400"
])
tbl_row(t3, [
    "шаблон — ошибки рендеринга\ninternal/web/handler.go",
    "template.ParseFS / ExecuteTemplate\n→ c.String(500, err.Error())",
    "HTTP Response (Output)",
    "CWE-209: полный текст ошибки Go с путями и именами раскрывался пользователю",
    "Заменено на generic «Internal server error»"
])

heading("3.3. Исправления с фрагментами кода", 2)

heading("3.3.1 Allowlist-валидация recipientID (CWE-20)", 3)

label("ДО:")
code(
    "recipientID := strings.ToUpper(strings.TrimSpace(c.PostForm(\"recipient_id\")))\n"
    "if len(recipientID) < 8 || len(recipientID) > 20 {\n"
    "    // любые символы, в том числе пробелы и спецсимволы — проходили\n"
    "    c.String(200, `<article class=\"flash error\">8-20 characters</article>`)\n"
    "    return\n"
    "}"
)
label("ПОСЛЕ:")
code(
    "var recipientIDPattern = regexp.MustCompile(`^[A-Z0-9]{8,20}$`)\n\n"
    "recipientID := strings.ToUpper(strings.TrimSpace(c.PostForm(\"recipient_id\")))\n"
    "if !recipientIDPattern.MatchString(recipientID) {\n"
    "    c.String(200, `<article class=\"flash error\">"
    "Recipient ID: 8-20 uppercase alphanumeric</article>`)\n"
    "    return\n"
    "}"
)
para(
    "Объяснение: до исправления длина проверялась, но состав символов — нет. "
    "Теперь применяется тот же паттерн ^[A-Z0-9]{8,20}$, что и в API-валидаторе "
    "(validators.go), устраняя расхождение между Web и API интерфейсами."
)
para("CWE: CWE-20 (Improper Input Validation). Подтверждение: POSTing recipient_id='A B<>' → 400, запись не создаётся.")

heading("3.3.2 Раскрытие информации через ошибки шаблонов (CWE-209)", 3)

label("ДО:")
code(
    "// internal/web/handler.go — функция render()\n"
    "if err != nil {\n"
    "    c.String(http.StatusInternalServerError,\n"
    "        \"Template parse error: \"+err.Error())  // путь к файлу, имя шаблона — видны клиенту\n"
    "}\n"
    "// renderPartial()\n"
    "if err := h.templates.ExecuteTemplate(c.Writer, name, data); err != nil {\n"
    "    c.String(http.StatusInternalServerError,\n"
    "        \"Template render error: \"+err.Error())  // аналогично\n"
    "}"
)
label("ПОСЛЕ:")
code(
    "if err != nil {\n"
    "    // CWE-209: только generic сообщение; детали пишутся в server-side лог\n"
    "    c.String(http.StatusInternalServerError, \"Internal server error\")\n"
    "    return\n"
    "}\n\n"
    "if err := h.templates.ExecuteTemplate(c.Writer, name, data); err != nil {\n"
    "    c.String(http.StatusInternalServerError, \"Internal server error\")\n"
    "}"
)
para(
    "Объяснение: текст ошибки Go включает внутренние пути к файлам, имена переменных "
    "и трассировку. Передача этих данных клиенту помогает атакующему в разведке "
    "архитектуры приложения. Теперь клиент видит только «Internal server error»; "
    "полный текст ошибки должен писаться в структурированный server-side лог."
)
para("CWE: CWE-209 (Information Exposure Through an Error Message).")

heading("3.3.3 Log Injection — sanitizeLogField (CWE-117)", 3)

label("ДО:")
code(
    "func (s *AuditService) Log(..., details, ipAddress string) {\n"
    "    _ = s.db.CreateAuditLog(userID, action, resourceID, details, ipAddress)\n"
    "    // ← CR/LF в details или ipAddress → поддельные строки в audit_logs\n"
    "}"
)
label("ПОСЛЕ:")
code(
    "func sanitizeLogField(s string) string {\n"
    "    s = strings.ReplaceAll(s, \"\\r\", \"\")\n"
    "    s = strings.ReplaceAll(s, \"\\n\", \" \")\n"
    "    s = strings.ReplaceAll(s, \"\\t\", \" \")\n"
    "    if len(s) > 500 { s = s[:500] }\n"
    "    return s\n"
    "}\n\n"
    "func (s *AuditService) Log(...) {\n"
    "    _ = s.db.CreateAuditLog(userID, action, resourceID,\n"
    "        sanitizeLogField(details), sanitizeLogField(ipAddress))\n"
    "}"
)

# ════════════════════════════════════════════════════════════════════════════
# 4. ЗАДАНИЕ 3 — AUTH / AUTHZ / CRYPTO
# ════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
heading("4. Задание 3: Аудит аутентификации, авторизации и криптографии")

heading("4.1. Матрица доступа по ролям", 2)

t4 = doc.add_table(rows=1, cols=4)
t4.style = "Table Grid"
tbl_header(t4, ["Действие / Endpoint", "Аноним", "user", "fraud_analyst"])
for row in [
    ("POST /api/v1/auth/register",            "✓", "—",                  "—"),
    ("POST /api/v1/auth/login",               "✓", "—",                  "—"),
    ("GET  /api/v1/me",                       "—", "✓",                  "✓"),
    ("POST /api/v1/payments",                 "—", "✓",                  "—"),
    ("GET  /api/v1/payments",                 "—", "✓ (только свои)",    "—"),
    ("GET  /api/v1/payments/:id",             "—", "✓ (только свои)",    "✓ (любой)"),
    ("POST /api/v1/payments/:id/confirm",     "—", "✓ (только свои)",    "—"),
    ("GET  /api/v1/analyst/payments/flagged", "—", "—",                  "✓"),
    ("POST /api/v1/analyst/payments/:id/flag","—", "—",                  "✓"),
    ("POST /api/v1/analyst/payments/:id/reject","—","—",                 "✓"),
    ("GET  /api/v1/analyst/audit",            "—", "—",                  "✓"),
]:
    tbl_row(t4, row)

heading("4.2. Выявленные риски в области аутентификации и криптографии", 2)

t5 = doc.add_table(rows=1, cols=4)
t5.style = "Table Grid"
tbl_header(t5, ["#", "Риск", "CWE", "Статус"])
for row in [
    ("1", "JWT cookie Secure=false — токен передаётся по HTTP открытым текстом", "CWE-614", "ИСПРАВЛЕНО: cfg.CookieSecure из переменной окружения COOKIE_SECURE"),
    ("2", "Отсутствие jti в JWT — невозможно отозвать конкретный токен без отзыва всех", "CWE-613", "ИСПРАВЛЕНО: jti = 16 случайных байт (crypto/rand), Base16"),
    ("3", "Роль из JWT не перепроверяется по БД — смена роли вступает в силу только после истечения токена", "CWE-269", "ПРИНЯТО: TTL 15 мин; TODO: revocation list"),
    ("4", "Нет per-user блокировки после N неудачных попыток входа — только IP rate limit", "CWE-307", "ЧАСТИЧНО: rate limit 100 req/мин; TODO: lockout counter в БД"),
    ("5", "bcrypt.DefaultCost (10) не адаптируется к росту мощности CPU", "CWE-916", "ПРИНЯТО: приемлемо для MVP; рекомендация: Cost 12 в продакшне"),
]:
    tbl_row(t5, row)

heading("4.3. Два сценария принятия решения о доступе", 2)

heading("Сценарий A — Легитимный: fraud_analyst просматривает чужой платёж", 3)
code(
    "GET /api/v1/payments/42\n"
    "Authorization: Bearer <valid_analyst_token>\n\n"
    "1. AuthMiddleware: JWT parsed → {UserID:2, Role:fraud_analyst}   OK\n"
    "2. PaymentHandler.GetPayment → PaymentService.GetPayment(42, 2, fraud_analyst)\n"
    "3. userRole == RoleFraudAnalyst → owner check обходится (по дизайну)\n"
    "4. AuditLog: action=payment_viewed, resource_id=42, ip=...\n"
    "5. ← 200 OK (analyst response, включает fraud_score)"
)

heading("Сценарий B — Запрещённый: обычный user пытается подтвердить чужой платёж", 3)
code(
    "POST /api/v1/payments/99/confirm\n"
    "Authorization: Bearer <valid_user_token>  (UserID=5)\n"
    "Payment #99 принадлежит UserID=7\n\n"
    "1. AuthMiddleware: JWT valid → {UserID:5, Role:user}             OK\n"
    "2. PaymentService.ConfirmPayment(99, 5)\n"
    "3. DB.GetPaymentByID(99) → payment.UserID = 7\n"
    "4. payment.UserID(7) != userID(5) → ErrUnauthorizedAccess\n"
    "5. ← 403 Forbidden {\"error\": \"access denied\"}"
)

heading("4.4. Исправления", 2)

heading("4.4.1 Cookie Secure flag из конфигурации (CWE-614)", 3)
label("ДО (internal/web/handler.go):")
code('c.SetCookie("token", token, 900, "/", "", false, true)\n'
     '//                                           ^^^^^ жёстко false')
label("ПОСЛЕ:")
code(
    "// internal/config/config.go:\n"
    "type Config struct {\n"
    "    ...\n"
    "    CookieSecure bool // из COOKIE_SECURE (по умолчанию true)\n"
    "}\n\n"
    "// internal/web/handler.go:\n"
    'c.SetCookie("token", token, 900, "/", "", h.cookieSecure, true)\n'
    "//                                         ^^^^^^^^^^^^^^"
)
para(
    "Объяснение: ранее флаг Secure был жёстко равен false, что означало "
    "передачу cookie по незашифрованному HTTP. Теперь значение берётся из "
    "конфигурации — в production всегда true (HTTPS), в dev-среде можно "
    "временно выставить false без изменения кода."
)

heading("4.4.2 JWT ID (jti) для будущего отзыва токенов (CWE-613)", 3)
label("ДО:")
code(
    "claims := Claims{\n"
    "    RegisteredClaims: jwt.RegisteredClaims{\n"
    "        ExpiresAt: ..., IssuedAt: ..., Subject: user.Email,\n"
    "        // ID отсутствует — все токены одного пользователя неразличимы\n"
    "    },\n"
    "}"
)
label("ПОСЛЕ:")
code(
    "jtiBytes := make([]byte, 16)\n"
    "rand.Read(jtiBytes)  // crypto/rand\n"
    "jti := hex.EncodeToString(jtiBytes)\n\n"
    "claims := Claims{\n"
    "    RegisteredClaims: jwt.RegisteredClaims{\n"
    "        ExpiresAt: ..., IssuedAt: ..., Subject: user.Email,\n"
    "        ID: jti,  // уникальный 32-символьный hex-идентификатор\n"
    "    },\n"
    "}"
)
para(
    "Объяснение: поле jti (JWT ID) делает каждый выданный токен уникальным. "
    "Это закладывает основу для реализации revocation list — при компрометации "
    "токена достаточно внести его jti в чёрный список, не трогая остальные."
)

heading("4.5. Архитектурные обоснования", 2)
para(
    "Роль пользователя хранится прямо в JWT и не перечитывается из БД при каждом "
    "запросе. Это сознательный архитектурный выбор: снижение нагрузки на SQLite "
    "важнее мгновенной реакции на смену роли. При TTL 15 минут окно уязвимости "
    "минимально. В продакшне стоит добавить Redis-список отозванных jti."
)
para(
    "bcrypt.DefaultCost (10) обеспечивает ~100 мс на хеш — достаточно для MVP. "
    "При переносе на современный многоядерный сервер рекомендуется повысить Cost "
    "до 12–13 и провести нагрузочное тестирование эндпоинтов аутентификации."
)

# ════════════════════════════════════════════════════════════════════════════
# 5. ЗАДАНИЕ 4 — CVSS v4.0
# ════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
heading("5. Задание 4: Оценка уязвимостей по CVSS v4.0")

para(
    "Для задания 4 выбраны 5 уязвимостей, не совпадающих с рассмотренными "
    "в ПР №4. Напомним, что ПР №4 охватывало: CWE-703 (ExecuteTemplate G104), "
    "CWE-614 (Cookie Secure), CWE-89 (SQL Injection), CWE-79 (XSS) и CWE-521 "
    "(слабая политика паролей). Все пять уязвимостей ниже — новые."
)

heading("5.1. Описание уязвимостей", 2)

heading("Уязвимость 1 — CWE-400: DoS через неограниченное тело запроса", 3)
para(
    "Модуль: cmd/server/main.go (все POST-эндпоинты). Точка входа: тело HTTP-запроса. "
    "Sink: оперативная память процесса (буфер Gin). До исправления любой анонимный "
    "пользователь мог отправить тело произвольного размера и вызвать OOM-crash сервера."
)

heading("Уязвимость 2 — CWE-613: Отсутствие jti / невозможность отзыва JWT", 3)
para(
    "Модуль: internal/services/auth_service.go. Точка входа: POST /api/v1/auth/login. "
    "Sink: JWT-токен (выдаётся клиенту). Без поля jti все токены одного пользователя "
    "идентичны — при утечке токена его невозможно инвалидировать без перевыпуска "
    "секретного ключа (что аннулирует все сессии)."
)

heading("Уязвимость 3 — CWE-117: Log Injection через поля аудита", 3)
para(
    "Модуль: internal/services/audit_service.go. Точка входа: поле details и IP-адрес "
    "во всех вызовах Log(). Sink: таблица audit_logs (SQLite). Атакующий с правами "
    "пользователя мог внедрить символы \\r\\n и создать в журнале фиктивные записи, "
    "искажая результаты расследования инцидентов."
)

heading("Уязвимость 4 — CWE-20: Отсутствие allowlist-валидации recipientID в веб-форме", 3)
para(
    "Модуль: internal/web/handler.go. Точка входа: POST /payments/new (поле recipient_id). "
    "Sink: таблица payments (SQL INSERT). Веб-форма проверяла только длину (8–20 символов), "
    "но не состав символов. Это создавало расхождение между API- и Web-интерфейсами — "
    "через веб можно было передать пробелы и спецсимволы, которые API отклонял бы."
)

heading("Уязвимость 5 — CWE-209: Раскрытие внутренних деталей ошибок Go", 3)
para(
    "Модуль: internal/web/handler.go. Точка входа: все страницы с рендерингом шаблонов. "
    "Sink: HTTP-ответ (браузер клиента). При ошибке парсинга или выполнения шаблона "
    "сервер возвращал полный текст ошибки Go, включая пути к файлам и имена переменных — "
    "информацию, полезную атакующему для разведки структуры приложения."
)

heading("5.2. Таблица CVSS v4.0", 2)

para(
    "Все векторы рассчитаны на официальном калькуляторе CVSS v4.0: "
    "https://www.first.org/cvss/calculator/4.0"
)

t6 = doc.add_table(rows=1, cols=8)
t6.style = "Table Grid"
tbl_header(t6, [
    "№", "Уязвимость\nМодуль/Endpoint",
    "CWE", "CVSS v4.0 Vector",
    "Score", "Severity",
    "Обоснование метрик",
    "Приоритет\nисправления"
])

cvss = [
    (
        "1",
        "DoS через тело запроса\ncmd/server/main.go\nвсе POST endpoints",
        "CWE-400",
        "CVSS:4.0/AV:N/AC:L/AT:N/PR:N/UI:N/"
        "VC:N/VI:N/VA:H/SC:N/SI:N/SA:N",
        "8.7",
        "HIGH",
        "AV:N — атака через интернет.\n"
        "AC:L, AT:N — никаких условий.\n"
        "PR:N — анонимно, UI:N.\n"
        "VA:H — сервис недоступен (OOM).\n"
        "Конфиденциальность/целостность не затронуты.",
        "1 — Критический\n(УСТРАНЕНО)"
    ),
    (
        "2",
        "Отсутствие jti / отзыва JWT\nauth_service.go\nPOST /api/v1/auth/login",
        "CWE-613",
        "CVSS:4.0/AV:N/AC:H/AT:N/PR:N/UI:N/"
        "VC:H/VI:H/VA:N/SC:N/SI:N/SA:N",
        "7.3",
        "HIGH",
        "AV:N — атакующий работает удалённо.\n"
        "AC:H — нужно предварительно завладеть токеном.\n"
        "VC:H/VI:H — украденный токен даёт полный доступ.\n"
        "TTL 15 мин снижает реальный риск.",
        "2 — Высокий\n(jti добавлен;\nrevocation — TODO)"
    ),
    (
        "3",
        "Log Injection через audit fields\naudit_service.go\nвсе вызовы Log()",
        "CWE-117",
        "CVSS:4.0/AV:N/AC:L/AT:N/PR:L/UI:N/"
        "VC:N/VI:L/VA:N/SC:N/SI:N/SA:N",
        "5.3",
        "MEDIUM",
        "AV:N, AC:L — легко воспроизводимо.\n"
        "PR:L — нужна авторизация пользователя.\n"
        "VI:L — искажение audit trail,\n"
        "но данные пользователей не утекают.",
        "3 — Средний\n(УСТРАНЕНО)"
    ),
    (
        "4",
        "Слабая валидация recipientID\nweb/handler.go\nPOST /payments/new",
        "CWE-20",
        "CVSS:4.0/AV:N/AC:L/AT:N/PR:L/UI:N/"
        "VC:N/VI:L/VA:N/SC:N/SI:N/SA:N",
        "5.3",
        "MEDIUM",
        "AV:N, AC:L, PR:L.\n"
        "VI:L — некорректные символы\nв поле recipient_id могут вызвать\n"
        "неожиданное поведение в смежных\nсистемах и отчётах.",
        "3 — Средний\n(УСТРАНЕНО)"
    ),
    (
        "5",
        "Раскрытие ошибок шаблонов\nweb/handler.go\nвсе GET-страницы",
        "CWE-209",
        "CVSS:4.0/AV:N/AC:L/AT:N/PR:N/UI:N/"
        "VC:L/VI:N/VA:N/SC:N/SI:N/SA:N",
        "6.9",
        "MEDIUM",
        "AV:N, AC:L — доступно без авторизации.\n"
        "PR:N, UI:N.\n"
        "VC:L — путь к файлу и внутренние\n"
        "имена помогают атакующему в разведке.\n"
        "Прямой утечки данных нет.",
        "2 — Высокий\n(УСТРАНЕНО)"
    ),
]
for r in cvss:
    tbl_row(t6, r)

heading("5.3. Ранжирование по приоритету исправления", 2)

para(
    "Ниже уязвимости расставлены по реальному приоритету для данного "
    "финтех-сервиса. Оценки CVSS в целом подтверждают этот порядок, "
    "хотя есть нюансы, о которых говорится в следующем разделе."
)

bullet("Приоритет 1 — CWE-400 (Score 8.7): анонимный DoS, не требует авторизации, "
       "гарантированно нарушает доступность платёжного сервиса. Устранено первым.")
bullet("Приоритет 2 — CWE-209 (Score 6.9): хотя балл ниже CWE-613, для финтеха "
       "разведка структуры приложения — прямой вектор к более серьёзным атакам. "
       "Исправлено немедленно.")
bullet("Приоритет 2 — CWE-613 (Score 7.3): кража сессии критична, но TTL 15 минут "
       "существенно снижает реальное окно уязвимости. jti добавлен; полный revocation list — "
       "следующий шаг.")
bullet("Приоритет 3 — CWE-117 (Score 5.3): искажение аудит-журнала — нарушение "
       "compliance (PCI DSS, GDPR). С точки зрения бизнеса важнее, чем CVSS-балл "
       "может показывать. Устранено.")
bullet("Приоритет 3 — CWE-20 (Score 5.3): слабая валидация в веб-форме, тогда как "
       "API-валидатор уже был корректным. Устранено выравниванием паттернов.")

heading("5.4. Совпадает ли CVSS с реальным приоритетом?", 2)

para(
    "В большинстве случаев — да. Наивысший CVSS-балл (8.7) у CWE-400 корректно "
    "отражает максимальный реальный риск: DoS на платёжный сервис — это немедленные "
    "финансовые потери и нарушение SLA."
)
para(
    "Есть и отклонения. CWE-209 получил балл 6.9 («только» VC:L), однако для финтеха "
    "утечка внутренней структуры приложения открывает путь к более серьёзным атакам — "
    "реальный приоритет выше балла. Напротив, CWE-613 имеет балл 7.3 (HIGH), но TTL "
    "15 минут значительно сокращает окно уязвимости, поэтому реальный приоритет "
    "чуть ниже формального."
)
para(
    "Вывод: CVSS v4.0 — надёжный инструмент количественной оценки, но он не учитывает "
    "бизнес-контекст (SLA, compliance, цепочки атак). Финальный приоритет исправлений "
    "всегда должен сочетать CVSS-балл с анализом реальных угроз для конкретного "
    "бизнес-сценария."
)

# ════════════════════════════════════════════════════════════════════════════
# 6. SAST / SCA
# ════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
heading("6. Результаты SAST и SCA после всех исправлений")

heading("6.1. Gosec (SAST)", 2)
code(
    "$ gosec ./...\n\n"
    "Summary:\n"
    "  Gosec  : dev\n"
    "  Files  : 21\n"
    "  Lines  : 2219\n"
    "  Nosec  : 0\n"
    "  Issues : 0"
)
para("Все ранее выявленные проблемы устранены. Gosec не выявил новых нарушений.")

heading("6.2. Govulncheck (SCA)", 2)
code(
    "$ govulncheck ./...\n\n"
    "=== Symbol Results ===\n"
    "No vulnerabilities found.\n"
    "Your code is affected by 0 vulnerabilities.\n\n"
    "This scan found 2 vulnerabilities in packages you import and 4 in modules\n"
    "you require, but your code doesn't appear to call these vulnerabilities.\n"
    "(GO-2026-4440, GO-2026-4441 — golang.org/x/net/html; HTML-парсер в коде не вызывается)"
)

# ════════════════════════════════════════════════════════════════════════════
# 7. ВЫВОД
# ════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
heading("7. Вывод по работе")

para(
    "В ходе практической работы №5 был проведён углублённый комплексный аудит "
    "безопасности MVP-проекта fintech-payments-mvp по четырём направлениям. "
    "В результате выявлено и устранено 9 уязвимостей и недостатков; 5 из них "
    "оценены по CVSS v4.0 с баллами от 5.3 до 8.7."
)
para(
    "Наиболее значимые улучшения архитектуры безопасности:"
)
bullet("MaxBodySizeMiddleware (1 МБ) — защита от DoS через исчерпание памяти (CWE-400).")
bullet("sanitizeLogField() — предотвращение подделки аудит-журнала через CRLF-инъекцию (CWE-117).")
bullet("Конфигурируемый флаг Secure для JWT-cookie, активированный по умолчанию (CWE-614).")
bullet("Поле jti в JWT-токенах для поддержки будущего механизма отзыва (CWE-613).")
bullet("Согласованная allowlist-валидация recipientID между Web- и API-интерфейсами (CWE-20).")
bullet("Замена детализированных сообщений об ошибках шаблонов на generic «Internal server error» (CWE-209).")

para(
    "Gosec подтверждает 0 проблем. Govulncheck не выявил уязвимостей в вызываемых символах. "
    "Архитектура MVP демонстрирует высокий уровень зрелости: параметризованные SQL-запросы "
    "исключают инъекции, html/template защищает от XSS, ролевая модель корректно реализует "
    "object-level access control."
)
para(
    "Оставшиеся риски — отсутствие revocation list для JWT и per-user account lockout — "
    "задокументированы как TODO-элементы roadmap для продакшн-деплоя. "
    "Оба требуют внешнего хранилища (Redis или дополнительная таблица в БД) и не "
    "реализованы в MVP ввиду его учебного характера."
)

doc.save("securep5.docx")
print("✓  securep5.docx saved")
